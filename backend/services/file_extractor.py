"""
文件内容提取服务 - 文件即知识库
支持 PDF、Word、Excel、CSV、TXT、JSON 等格式
"""
import json
import csv
import io
import chardet
from pathlib import Path
from typing import Optional
from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook

from core.logging import get_logger

logger = get_logger(__name__)


class FileExtractor:
    """文件内容提取器 - 将各种文件转换为文本"""
    
    # 文件大小限制（字符数），超过则截断
    MAX_TEXT_LENGTH = 100000  # 约 10 万字符，可根据 LLM 上下文调整
    
    @staticmethod
    def extract(file_path: str, mime_type: Optional[str] = None) -> dict:
        """
        提取文件内容
        
        Returns:
            {
                "success": bool,
                "text": str,           # 提取的文本内容
                "truncated": bool,     # 是否被截断
                "original_size": int,  # 原始字符数
                "error": str,          # 错误信息（如果有）
            }
        """
        path = Path(file_path)
        if not path.exists():
            return {"success": False, "text": "", "truncated": False, "original_size": 0, "error": "文件不存在"}
        
        ext = path.suffix.lower()
        
        try:
            if ext == ".pdf":
                result = FileExtractor._extract_pdf(path)
            elif ext in [".docx", ".doc"]:
                result = FileExtractor._extract_word(path)
            elif ext in [".xlsx", ".xls"]:
                result = FileExtractor._extract_excel(path)
            elif ext == ".csv":
                result = FileExtractor._extract_csv(path)
            elif ext == ".json":
                result = FileExtractor._extract_json(path)
            elif ext in [".txt", ".md", ".py", ".js", ".html", ".css", ".xml"]:
                result = FileExtractor._extract_text(path)
            else:
                # 尝试作为文本读取
                result = FileExtractor._extract_text(path)
            
            # 截断处理
            original_size = len(result["text"])
            truncated = False
            text = result["text"]
            
            if original_size > FileExtractor.MAX_TEXT_LENGTH:
                text = text[:FileExtractor.MAX_TEXT_LENGTH]
                truncated = True
                logger.warning(f"文件内容过长已截断: {path.name} ({original_size} -> {FileExtractor.MAX_TEXT_LENGTH})")
            
            return {
                "success": result.get("success", True),
                "text": text,
                "truncated": truncated,
                "original_size": original_size,
                "error": result.get("error", ""),
            }
            
        except Exception as e:
            logger.error(f"文件提取失败 {path}: {e}")
            return {
                "success": False,
                "text": "",
                "truncated": False,
                "original_size": 0,
                "error": str(e),
            }
    
    @staticmethod
    def _extract_pdf(path: Path) -> dict:
        """提取 PDF 文本"""
        reader = PdfReader(str(path))
        text_parts = []
        
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"--- 第 {i + 1} 页 ---\n{page_text}")
            except Exception as e:
                text_parts.append(f"--- 第 {i + 1} 页 [无法提取] ---")
        
        return {
            "success": True,
            "text": "\n\n".join(text_parts),
        }
    
    @staticmethod
    def _extract_word(path: Path) -> dict:
        """提取 Word 文档文本"""
        doc = Document(str(path))
        text_parts = []
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # 提取表格
        for i, table in enumerate(doc.tables):
            text_parts.append(f"\n--- 表格 {i + 1} ---")
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text_parts.append(row_text)
        
        return {
            "success": True,
            "text": "\n".join(text_parts),
        }
    
    @staticmethod
    def _extract_excel(path: Path) -> dict:
        """提取 Excel 表格文本"""
        wb = load_workbook(str(path), data_only=True)
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"=== 工作表: {sheet_name} ===")
            
            # 提取所有行
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(
                    str(cell) if cell is not None else "" 
                    for cell in row
                )
                if row_text.strip():
                    text_parts.append(row_text)
            
            text_parts.append("")  # 空行分隔
        
        return {
            "success": True,
            "text": "\n".join(text_parts),
        }
    
    @staticmethod
    def _extract_csv(path: Path) -> dict:
        """提取 CSV 文本"""
        # 检测编码
        with open(path, "rb") as f:
            raw_data = f.read()
            encoding = chardet.detect(raw_data)["encoding"] or "utf-8"
        
        text_parts = []
        with open(path, "r", encoding=encoding, errors="replace") as f:
            reader = csv.reader(f)
            for row in reader:
                row_text = " | ".join(row)
                text_parts.append(row_text)
        
        return {
            "success": True,
            "text": "\n".join(text_parts),
        }
    
    @staticmethod
    def _extract_json(path: Path) -> dict:
        """提取并格式化 JSON"""
        with open(path, "rb") as f:
            raw_data = f.read()
            encoding = chardet.detect(raw_data)["encoding"] or "utf-8"
        
        with open(path, "r", encoding=encoding, errors="replace") as f:
            data = json.load(f)
        
        # 格式化为易读的 JSON
        formatted = json.dumps(data, ensure_ascii=False, indent=2)
        
        return {
            "success": True,
            "text": formatted,
        }
    
    @staticmethod
    def _extract_text(path: Path) -> dict:
        """提取纯文本文件"""
        with open(path, "rb") as f:
            raw_data = f.read()
            encoding = chardet.detect(raw_data)["encoding"] or "utf-8"
        
        with open(path, "r", encoding=encoding, errors="replace") as f:
            text = f.read()
        
        return {
            "success": True,
            "text": text,
        }


# 便捷函数
def extract_file_content(file_path: str, mime_type: Optional[str] = None) -> dict:
    """提取文件内容的便捷函数"""
    return FileExtractor.extract(file_path, mime_type)
